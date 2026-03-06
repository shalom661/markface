from docx import Document
import sys

path = r"c:\MARK FACE HUB\MATÉRIA PRIMA.docx"
try:
    doc = Document(path)
    for p in doc.paragraphs:
        if p.text.strip():
            print(p.text)
    print("\n--- TABLES ---")
    for i, t in enumerate(doc.tables):
        print(f"Table {i+1}:")
        for row in t.rows:
            print(" | ".join([c.text.strip().replace('\n', ' ') for c in row.cells]))
        print("-" * 20)
except Exception as e:
    print(f"Error: {e}")
